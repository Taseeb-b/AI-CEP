# -*- coding: utf-8 -*-
"""
==========================================================================================
AI-POWERED AUTONOMOUS LITERATURE REVIEW AGENT (Agentic RAG-lite System) - Streamlit version
Plans search strategies, autonomously refines queries based on coverage, retrieves and
analyzes arXiv papers, clusters themes, detects gaps and contradictions, and generates a
professionally formatted Word literature review with proper citations.

DEPLOYMENT: Streamlit Community Cloud
  - requirements.txt must list: streamlit, openai, arxiv, pymupdf, python-docx, pandas,
    requests, tenacity
  - Enter your Groq API key inside the app itself (get one at https://console.groq.com/keys)
==========================================================================================
"""

import os
import re
import json
import time
import traceback
from dataclasses import dataclass, field
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple

import requests
import pandas as pd
import streamlit as st
import arxiv
import fitz  # PyMuPDF

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from openai import OpenAI
from tenacity import retry, stop_after_attempt, wait_exponential

MODEL_NAME_DEFAULT = "llama-3.1-8b-instant"  # far higher daily/request quota on Groq's
                                              # free tier than the 70b model - better fit
                                              # for a pipeline that makes many calls
GROQ_BASE_URL = "https://api.groq.com/openai/v1"

# =========================================================================================
# LLM client wrapper
# =========================================================================================

class LLMClient:
    """Thin wrapper around Groq's API (OpenAI-compatible), used for every reasoning step.

    Groq's free tier caps BOTH requests-per-minute AND tokens-per-minute (e.g. 30 RPM / 12,000
    TPM for the 70b model). A fixed delay between calls only protects the RPM side - a handful
    of large summarization calls can still blow the TPM budget in seconds. This client tracks
    actual token usage in a rolling 60-second window and sleeps just long enough to stay under
    both limits before every call, so the pipeline self-paces instead of erroring out.
    """

    def __init__(self, api_key: str, model: str = MODEL_NAME_DEFAULT,
                 max_requests_per_min: int = 25, max_tokens_per_min: int = 5500):
        if not api_key:
            raise ValueError("A Groq API key is required.")
        self.client = OpenAI(api_key=api_key, base_url=GROQ_BASE_URL)
        self.model = model
        self.max_requests_per_min = max_requests_per_min
        self.max_tokens_per_min = max_tokens_per_min
        self._history: List[Tuple[float, int]] = []  # (timestamp, tokens_used) in the last 60s

    @staticmethod
    def _estimate_tokens(*texts: str) -> int:
        # Rough, conservative heuristic: ~4 characters per token.
        return sum(len(t) for t in texts if t) // 4

    def _throttle(self, projected_tokens: int) -> None:
        now = time.time()
        self._history = [(ts, tok) for ts, tok in self._history if now - ts < 60]

        while self._history:
            req_count = len(self._history)
            tok_count = sum(tok for _, tok in self._history) + projected_tokens
            if req_count < self.max_requests_per_min and tok_count <= self.max_tokens_per_min:
                break
            oldest_ts = self._history[0][0]
            sleep_for = max(0.5, 60 - (now - oldest_ts) + 0.5)
            print(f"[Agent] Approaching Groq rate limit - pausing {sleep_for:.0f}s to stay within quota...")
            time.sleep(sleep_for)
            now = time.time()
            self._history = [(ts, tok) for ts, tok in self._history if now - ts < 60]

    @retry(stop=stop_after_attempt(6), wait=wait_exponential(multiplier=2, min=4, max=60))
    def complete(self, system: str, user: str, max_tokens: int = 1500) -> str:
        projected = self._estimate_tokens(system, user) + max_tokens
        self._throttle(projected)
        resp = None
        try:
            resp = self.client.chat.completions.create(
                model=self.model,
                max_tokens=max_tokens,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
            )
        finally:
            actual = getattr(getattr(resp, "usage", None), "total_tokens", None) if resp is not None else None
            self._history.append((time.time(), actual or projected))
        return resp.choices[0].message.content or ""

    def complete_json(self, system: str, user: str, max_tokens: int = 1500) -> Any:
        full_system = system + "\n\nRespond ONLY with valid JSON. No prose, no markdown fences."
        raw = self.complete(full_system, user, max_tokens=max_tokens)
        cleaned = re.sub(r"^```json|```$", "", raw.strip(), flags=re.MULTILINE).strip()
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            match = re.search(r"(\{.*\}|\[.*\])", cleaned, re.DOTALL)
            if match:
                try:
                    return json.loads(match.group(0))
                except json.JSONDecodeError:
                    pass
            return {"error": "Could not parse model output as JSON", "raw": raw}


# =========================================================================================
# Config & data model
# =========================================================================================

@dataclass
class ReviewConfig:
    topic: str = ""
    num_papers: int = 10
    depth: str = "Standard"          # Quick | Standard | Comprehensive
    length: str = "5-10 pages"       # 3-5 pages | 5-10 pages | 10-20 pages | Custom
    custom_words: int = 3000
    citation_style: str = "IEEE"     # IEEE | APA | MLA | Chicago


@dataclass
class PaperRecord:
    arxiv_id: str
    title: str
    authors: List[str]
    year: str
    abstract: str
    pdf_url: str
    citation_key: str = ""          # e.g. "P1" - internal tag used by the LLM when citing
    relevance_score: float = 0.0
    relevance_reason: str = ""
    text_excerpt: str = ""          # extracted PDF text (truncated)
    pdf_ok: bool = False
    summary: Dict[str, Any] = field(default_factory=dict)   # structured comparative fields
    themes: List[str] = field(default_factory=list)

    def author_display(self) -> str:
        if not self.authors:
            return "Unknown"
        first = self.authors[0].strip().split()
        surname = first[-1] if first else self.authors[0]
        if len(self.authors) == 1:
            return surname
        elif len(self.authors) == 2:
            second = self.authors[1].strip().split()
            surname2 = second[-1] if second else self.authors[1]
            return f"{surname} & {surname2}"
        else:
            return f"{surname} et al."


DEPTH_ROUNDS = {"Quick": 1, "Standard": 2, "Comprehensive": 3}
LENGTH_WORDS = {"3-5 pages": 2000, "5-10 pages": 3500, "10-20 pages": 7000}


def target_word_count(cfg: ReviewConfig) -> int:
    if cfg.length == "Custom":
        return max(800, cfg.custom_words)
    return LENGTH_WORDS.get(cfg.length, 3500)


# =========================================================================================
# Phase 1: Research planning / query generation
# =========================================================================================

QUERY_GEN_SYSTEM = """You are a research planning agent. Given a research topic, generate an
effective set of arXiv search queries that will maximize both precision and coverage. Include:
the core topic phrased naturally, 2-3 synonyms/related terminology variants, one broader query,
and one narrower/more specific query. Return a JSON object with key "queries": a list of 4-6
short search query strings (each 2-6 words, suitable for a search engine, no boolean operators)."""


def generate_search_queries(llm: LLMClient, topic: str) -> List[str]:
    result = llm.complete_json(QUERY_GEN_SYSTEM, f"Research topic: {topic}", max_tokens=400)
    queries = result.get("queries") if isinstance(result, dict) else None
    if not queries:
        queries = [topic]
    queries = [q.strip() for q in queries if q and q.strip()]
    if topic not in queries:
        queries.insert(0, topic)
    return queries[:6]


# =========================================================================================
# Phase 2: Autonomous literature search (arXiv)
# =========================================================================================

RECENCY_YEARS = 5  # only keep papers published within the last N years (rolling, not hardcoded)


def search_arxiv(queries: List[str], max_results_per_query: int = 12,
                  min_year: Optional[int] = None) -> List[PaperRecord]:
    """Searches arXiv for each query and returns a de-duplicated list of PaperRecords.
    Uses arXiv's official API (stable, no key required) instead of scraping.
    Papers older than `min_year` (default: current year - RECENCY_YEARS) are dropped, so the
    review stays focused on recent literature instead of pulling in decade-old papers."""
    if min_year is None:
        min_year = datetime.now().year - RECENCY_YEARS

    client = arxiv.Client(page_size=max_results_per_query, delay_seconds=3, num_retries=3)
    seen_ids, papers = set(), []
    skipped_old = 0

    for q in queries:
        try:
            search = arxiv.Search(query=q, max_results=max_results_per_query,
                                   sort_by=arxiv.SortCriterion.Relevance)
            for r in client.results(search):
                arxiv_id = r.get_short_id()
                if arxiv_id in seen_ids:
                    continue
                seen_ids.add(arxiv_id)
                year = r.published.year if r.published else None
                if year is not None and year < min_year:
                    skipped_old += 1
                    continue
                papers.append(PaperRecord(
                    arxiv_id=arxiv_id,
                    title=r.title.strip().replace(chr(10), " "),
                    authors=[a.name for a in r.authors],
                    year=str(year) if year else "n.d.",
                    abstract=r.summary.strip().replace(chr(10), " "),
                    pdf_url=r.pdf_url,
                ))
        except Exception as e:
            print(f"[Agent] arXiv search failed for query '{q}': {e}")
            continue

    if skipped_old:
        print(f"[Agent] Filtered out {skipped_old} paper(s) older than {min_year}.")

    return papers


# =========================================================================================
# Phase 3: Relevance evaluation
# =========================================================================================

RELEVANCE_SYSTEM = """You are a research relevance evaluator. Given a research topic and a batch
of candidate papers (title + abstract), score each paper's relevance to the topic from 0-100,
considering: direct relevance of subject matter, abstract quality/specificity, and whether it
addresses the topic's methods or application domain. Return a JSON object with key "scores": a
list of objects {id, score, reason} where id matches the paper id given, score is 0-100, and
reason is a one-sentence justification."""


def score_relevance(llm: LLMClient, topic: str, papers: List[PaperRecord], batch_size: int = 10) -> None:
    """Scores papers in batches (mutates PaperRecord.relevance_score/relevance_reason in place)."""
    for i in range(0, len(papers), batch_size):
        batch = papers[i:i + batch_size]
        listing = "\n\n".join(
            f"id: {p.arxiv_id}\ntitle: {p.title}\nabstract: {p.abstract[:600]}"
            for p in batch
        )
        user_prompt = f"Research topic: {topic}\n\nCandidate papers:\n\n{listing}"
        result = llm.complete_json(RELEVANCE_SYSTEM, user_prompt, max_tokens=1200)
        scores = result.get("scores") if isinstance(result, dict) else None
        score_map = {}
        if scores:
            for s in scores:
                try:
                    score_map[s["id"]] = (float(s.get("score", 0)), s.get("reason", ""))
                except Exception:
                    continue
        for p in batch:
            sc, reason = score_map.get(p.arxiv_id, (0.0, "Not scored."))
            p.relevance_score = sc
            p.relevance_reason = reason


# =========================================================================================
# Phase 4: Coverage analysis & autonomous replanning
# =========================================================================================

COVERAGE_SYSTEM = """You are a coverage analyst for a literature review agent. Given a research
topic and the titles/abstracts of the top candidate papers found so far, decide whether these
papers sufficiently cover the major subtopics, methods, and application angles of the research
topic. If coverage is insufficient, identify the specific missing subtopics and propose 2-4 new,
more targeted arXiv search queries to fill those gaps (do not repeat prior queries).
Return strict JSON: {"sufficient": bool, "missing_subtopics": [strings], "new_queries": [strings]}."""


def analyze_coverage(llm: LLMClient, topic: str, top_papers: List[PaperRecord],
                      prior_queries: List[str]) -> Dict[str, Any]:
    listing = "\n\n".join(f"- {p.title}: {p.abstract[:300]}" for p in top_papers[:15])
    user_prompt = (f"Research topic: {topic}\n\nQueries already tried: {prior_queries}\n\n"
                   f"Top papers found so far:\n\n{listing}")
    result = llm.complete_json(COVERAGE_SYSTEM, user_prompt, max_tokens=700)
    if not isinstance(result, dict):
        result = {"sufficient": True, "missing_subtopics": [], "new_queries": []}
    return result


def run_search_and_select(llm: LLMClient, cfg: ReviewConfig, progress=None) -> Tuple[List[PaperRecord], List[str]]:
    """Core autonomous loop: plan queries -> search -> score -> check coverage -> replan if
    needed, up to a depth-controlled number of rounds. Returns the top-N selected papers."""
    max_rounds = DEPTH_ROUNDS.get(cfg.depth, 2)
    all_papers: Dict[str, PaperRecord] = {}
    tried_queries: List[str] = []

    if progress:
        progress(0.05, desc="Planning search strategy...")
    queries = generate_search_queries(llm, cfg.topic)

    for round_num in range(1, max_rounds + 1):
        if progress:
            progress(0.05 + 0.15 * round_num / max_rounds,
                      desc=f"Searching literature (round {round_num}/{max_rounds})...")
        new_papers = search_arxiv(queries, max_results_per_query=8)  # lowered to bound total pool size
        tried_queries.extend(queries)
        added = 0
        for p in new_papers:
            if p.arxiv_id not in all_papers:
                all_papers[p.arxiv_id] = p
                added += 1
        print(f"[Agent] Round {round_num}: {added} new paper(s), {len(all_papers)} total.")

        if progress:
            progress(0.25, desc="Evaluating relevance...")
        # Only score papers we haven't scored yet - re-scoring everything found in earlier
        # rounds every round wastes calls/tokens and was a major cause of slow, rate-limited runs.
        # Also hard-cap how many NEW papers get scored per round regardless of how many arXiv
        # returned - without this, a broad topic across multiple search rounds can balloon into
        # 70-100+ papers and many extra minutes of unavoidable free-tier rate-limit waiting.
        CANDIDATE_CAP_PER_ROUND = 25
        unscored = [p for p in all_papers.values() if not p.relevance_reason][:CANDIDATE_CAP_PER_ROUND]
        if unscored:
            score_relevance(llm, cfg.topic, unscored)

        ranked = sorted(all_papers.values(), key=lambda p: p.relevance_score, reverse=True)
        top_slice = ranked[: max(cfg.num_papers, 10)]

        if progress:
            progress(0.32, desc="Checking topic coverage...")
        coverage = analyze_coverage(llm, cfg.topic, top_slice, tried_queries)

        if coverage.get("sufficient", True) or round_num == max_rounds:
            break

        queries = coverage.get("new_queries") or []
        if not queries:
            break
        print(f"[Agent] Coverage insufficient - missing: {coverage.get('missing_subtopics')}. "
              f"Replanning with: {queries}")

    ranked = sorted(all_papers.values(), key=lambda p: p.relevance_score, reverse=True)
    selected = [p for p in ranked if p.relevance_score > 0][: cfg.num_papers]
    if not selected:
        selected = ranked[: cfg.num_papers]

    for i, p in enumerate(selected):
        p.citation_key = f"P{i + 1}"

    return selected, tried_queries


# =========================================================================================
# Phase 5-6: PDF retrieval & text extraction
# =========================================================================================

def download_and_extract(papers: List[PaperRecord], workdir: str = "/tmp/litreview_pdfs",
                          progress=None) -> None:
    """Downloads each paper's PDF and extracts text via PyMuPDF. On any failure, the paper is
    kept but marked pdf_ok=False and the abstract alone is used as its evidence base."""
    os.makedirs(workdir, exist_ok=True)
    for i, p in enumerate(papers):
        if progress:
            progress(0.4 + 0.2 * (i + 1) / max(len(papers), 1),
                      desc=f"Downloading & extracting: {p.title[:45]}")
        try:
            resp = requests.get(p.pdf_url, timeout=20, headers={"User-Agent": "Mozilla/5.0"})
            resp.raise_for_status()
            path = os.path.join(workdir, f"{p.arxiv_id.replace('/', '_')}.pdf")
            with open(path, "wb") as f:
                f.write(resp.content)

            doc = fitz.open(path)
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            text = re.sub(r"\n{3,}", "\n\n", text).strip()

            if len(text) < 200:
                raise ValueError("Extracted text too short - likely a corrupted/unreadable PDF.")

            # Keep a head (intro/methodology) and tail (results/conclusion) slice within a
            # token-friendly budget rather than the full paper.
            head, tail = text[:5000], text[-3000:]
            p.text_excerpt = head + "\n\n[...]\n\n" + tail
            p.pdf_ok = True
        except Exception as e:
            print(f"[Agent] PDF failed for {p.arxiv_id}: {e}")
            p.text_excerpt = p.abstract
            p.pdf_ok = False


# =========================================================================================
# Phase 6b: Structured paper summarization (evidence grounding)
# =========================================================================================

SUMMARY_SYSTEM = """You are a research paper analyst. Given the extracted text (or abstract-only
if the PDF was unavailable) of one paper, extract a structured summary. Be concise and factual -
only state what the text supports; if a field is not discussed, use "Not specified". Return
strict JSON with keys: dataset (string), methodology (string), algorithms (string),
evaluation_metrics (string), key_results (string), strengths (string), limitations (string),
main_finding (1-2 sentence string)."""


def summarize_paper(llm: LLMClient, topic: str, paper: PaperRecord) -> None:
    user_prompt = (f"Research topic context: {topic}\n\nPaper title: {paper.title}\n\n"
                   f"Extracted content:\n{paper.text_excerpt[:4000]}")
    result = llm.complete_json(SUMMARY_SYSTEM, user_prompt, max_tokens=500)
    if isinstance(result, dict):
        paper.summary = result


def summarize_all_papers(llm: LLMClient, cfg: ReviewConfig, papers: List[PaperRecord], progress=None) -> None:
    for i, p in enumerate(papers):
        if progress:
            progress(0.6 + 0.1 * (i + 1) / max(len(papers), 1),
                      desc=f"Analyzing: {p.title[:45]}")
        try:
            summarize_paper(llm, cfg.topic, p)
        except Exception as e:
            print(f"[Agent] Summarization failed for {p.arxiv_id}: {e}")
            p.summary = {"main_finding": "Summarization failed."}


# =========================================================================================
# Phase 8: Theme discovery
# =========================================================================================

THEME_SYSTEM = """You are a research synthesis agent. Given structured summaries of several papers
(each tagged with a citation key like P1, P2, ...), group them into 3-6 coherent research themes
(e.g. methodological families, application domains, or approaches). A paper may belong to more
than one theme if genuinely relevant to both. Return strict JSON: {"themes": [{"name": string,
"description": string (1 sentence), "paper_keys": [string, ...]}]}."""


def discover_themes(llm: LLMClient, topic: str, papers: List[PaperRecord]) -> List[Dict[str, Any]]:
    listing = "\n\n".join(
        f"[{p.citation_key}] {p.title}\nMethodology: {p.summary.get('methodology','n/a')}\n"
        f"Main finding: {p.summary.get('main_finding','n/a')}"
        for p in papers
    )
    result = llm.complete_json(THEME_SYSTEM, f"Research topic: {topic}\n\nPapers:\n\n{listing}", max_tokens=1200)
    themes = result.get("themes") if isinstance(result, dict) else []
    themes = themes or []

    key_to_paper = {p.citation_key: p for p in papers}
    for t in themes:
        for key in t.get("paper_keys", []):
            if key in key_to_paper:
                key_to_paper[key].themes.append(t.get("name", "Uncategorized"))
    return themes


# =========================================================================================
# Phase 9: Comparative analysis table
# =========================================================================================

def build_comparison_table(papers: List[PaperRecord]) -> pd.DataFrame:
    rows = []
    for p in papers:
        s = p.summary or {}
        rows.append({
            "Key": p.citation_key,
            "Title": p.title,
            "Authors": p.author_display(),
            "Year": p.year,
            "Dataset": s.get("dataset", "Not specified"),
            "Methodology": s.get("methodology", "Not specified"),
            "Algorithms": s.get("algorithms", "Not specified"),
            "Metrics": s.get("evaluation_metrics", "Not specified"),
            "Key Results": s.get("key_results", "Not specified"),
            "Strengths": s.get("strengths", "Not specified"),
            "Limitations": s.get("limitations", "Not specified"),
        })
    return pd.DataFrame(rows)


# =========================================================================================
# Phase 10-11: Research gap & contradiction detection
# =========================================================================================

GAP_SYSTEM = """You are a research gap analyst. Given structured summaries of the reviewed papers
and their themes, identify concrete, evidence-based research gaps - e.g. underexplored datasets,
rarely used algorithms, missing evaluation methods, geographic limitations, or unaddressed
application domains. Each gap MUST reference which papers/themes support the observation using
their citation keys. Do not invent gaps unsupported by the given summaries.
Return strict JSON: {"gaps": [{"gap": string, "evidence_keys": [string,...]}]}."""

CONTRADICTION_SYSTEM = """You are a contradiction analyst. Given structured summaries of the
reviewed papers (key_results, main_finding), identify pairs or groups of papers whose findings
or conclusions appear to conflict (e.g. one paper claims method A outperforms B, another claims
the opposite). Only report genuine conflicts you can support with the given text - if there are
none, return an empty list. Return strict JSON: {"contradictions": [{"description": string,
"paper_keys": [string,...], "possible_reason": string}]}."""


def detect_gaps(llm: LLMClient, topic: str, papers: List[PaperRecord], themes: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    listing = "\n\n".join(
        f"[{p.citation_key}] {p.title} - Themes: {p.themes} - Limitations: {p.summary.get('limitations','n/a')} "
        f"- Dataset: {p.summary.get('dataset','n/a')}"
        for p in papers
    )
    result = llm.complete_json(GAP_SYSTEM, f"Topic: {topic}\n\nThemes: {[t.get('name') for t in themes]}\n\nPapers:\n\n{listing}", max_tokens=1000)
    return (result.get("gaps") if isinstance(result, dict) else []) or []


def detect_contradictions(llm: LLMClient, papers: List[PaperRecord]) -> List[Dict[str, Any]]:
    listing = "\n\n".join(
        f"[{p.citation_key}] Main finding: {p.summary.get('main_finding','n/a')} - "
        f"Key results: {p.summary.get('key_results','n/a')}"
        for p in papers
    )
    result = llm.complete_json(CONTRADICTION_SYSTEM, f"Papers:\n\n{listing}", max_tokens=900)
    return (result.get("contradictions") if isinstance(result, dict) else []) or []


# =========================================================================================
# Citation formatting
# =========================================================================================

def format_reference_entry(p: PaperRecord, style: str, ieee_number: Optional[int] = None) -> str:
    authors = ", ".join(p.authors[:6]) + (" et al." if len(p.authors) > 6 else "")
    url = f"https://arxiv.org/abs/{p.arxiv_id}"
    if style == "IEEE":
        n = f"[{ieee_number}] " if ieee_number else ""
        return f'{n}{authors}, "{p.title}," arXiv:{p.arxiv_id}, {p.year}. [Online]. Available: {url}'
    elif style == "APA":
        return f"{authors} ({p.year}). {p.title}. arXiv:{p.arxiv_id}. {url}"
    elif style == "MLA":
        return f'{authors}. "{p.title}." arXiv, {p.year}, {url}.'
    else:  # Chicago
        return f'{authors}. "{p.title}." arXiv preprint arXiv:{p.arxiv_id} ({p.year}). {url}.'


# =========================================================================================
# Phase 12: Section-by-section literature review generation
# =========================================================================================

SECTION_SYSTEM = """You are an academic writer producing one section of a formal literature
review. Write in a professional, academic, evidence-based tone. Ground every claim in the
provided paper summaries - do not invent facts, datasets, or results not present in the evidence.
When referencing a specific paper's finding, cite it using its bracket tag exactly as given
(e.g. "...as shown in [P3]..." or "[P1][P4] both report..."). Do not fabricate a references list
or repeat the bracket-tag legend - just write flowing academic prose using the tags inline.
Do not use markdown headers (#) - the heading is added separately. Return plain prose paragraphs."""


def _evidence_block(papers: List[PaperRecord]) -> str:
    return "\n\n".join(
        f"[{p.citation_key}] \"{p.title}\" ({p.author_display()}, {p.year})\n"
        f"Main finding: {p.summary.get('main_finding','n/a')}\n"
        f"Methodology: {p.summary.get('methodology','n/a')} | Dataset: {p.summary.get('dataset','n/a')}\n"
        f"Results: {p.summary.get('key_results','n/a')}\n"
        f"Strengths: {p.summary.get('strengths','n/a')} | Limitations: {p.summary.get('limitations','n/a')}"
        for p in papers
    )


def generate_section(llm: LLMClient, section_name: str, instructions: str, topic: str,
                      papers: List[PaperRecord], target_words: int) -> str:
    evidence = _evidence_block(papers)
    user_prompt = (f"Research topic: {topic}\nSection to write: {section_name}\n"
                   f"Specific instructions: {instructions}\nTarget length: ~{target_words} words.\n\n"
                   f"Evidence base (cite using these bracket tags):\n\n{evidence}")
    max_tok = min(2200, max(400, target_words * 2))
    return llm.complete(SECTION_SYSTEM, user_prompt, max_tokens=max_tok)


def generate_theme_section(llm: LLMClient, theme: Dict[str, Any], topic: str,
                            papers_by_key: Dict[str, PaperRecord], target_words: int) -> str:
    theme_papers = [papers_by_key[k] for k in theme.get("paper_keys", []) if k in papers_by_key]
    if not theme_papers:
        return ""
    instructions = (f"Discuss this specific research theme: '{theme.get('name')}' - "
                    f"{theme.get('description','')}. Compare and synthesize the papers below, "
                    f"noting agreements, differences in approach, and relative performance where reported.")
    return generate_section(llm, theme.get("name", "Theme"), instructions, topic, theme_papers, target_words)


def build_full_review_text(llm: LLMClient, cfg: ReviewConfig, papers: List[PaperRecord],
                            themes: List[Dict[str, Any]], gaps: List[Dict[str, Any]],
                            contradictions: List[Dict[str, Any]], progress=None) -> Dict[str, str]:
    """Generates every section of the review and returns an ordered dict {heading: body_text}."""
    total_words = target_word_count(cfg)
    n_sections = 5 + len(themes)  # intro, background, N themes, comparative, gaps, conclusion
    per_section = max(250, total_words // n_sections)
    papers_by_key = {p.citation_key: p for p in papers}

    sections: Dict[str, str] = {}
    steps = ["Introduction", "Background"] + [t.get("name", f"Theme {i+1}") for i, t in enumerate(themes)] + \
            ["Comparative Discussion", "Research Gaps and Challenges", "Future Research Directions", "Conclusion"]

    for i, step in enumerate(steps):
        if progress:
            progress(0.75 + 0.2 * (i + 1) / len(steps), desc=f"Writing: {step}")

        if step == "Introduction":
            instr = ("Introduce the research topic, its importance/motivation, and state the "
                     "scope and objective of this literature review.")
            sections[step] = generate_section(llm, step, instr, cfg.topic, papers, per_section)
        elif step == "Background":
            instr = "Provide background/context a reader needs to understand the topic before diving into themes."
            sections[step] = generate_section(llm, step, instr, cfg.topic, papers, per_section)
        elif step == "Comparative Discussion":
            instr = ("Synthesize the comparison across all reviewed papers - highlight overall "
                     "trends in methodology, datasets, and performance across the literature.")
            sections[step] = generate_section(llm, step, instr, cfg.topic, papers, per_section)
        elif step == "Research Gaps and Challenges":
            gap_lines = "\n".join(f"- {g.get('gap')} (supported by {g.get('evidence_keys')})" for g in gaps)
            contra_lines = "\n".join(
                f"- {c.get('description')} (papers: {c.get('paper_keys')}, possible reason: {c.get('possible_reason')})"
                for c in contradictions
            )
            instr = (f"Discuss the following identified research gaps, weaving in citations:\n{gap_lines}\n\n"
                     f"Also discuss these conflicting findings across the literature:\n{contra_lines or 'None identified.'}")
            sections[step] = generate_section(llm, step, instr, cfg.topic, papers, per_section)
        elif step == "Future Research Directions":
            instr = "Based on the identified gaps and limitations, propose concrete future research directions."
            sections[step] = generate_section(llm, step, instr, cfg.topic, papers, per_section)
        elif step == "Conclusion":
            instr = "Summarize the key takeaways of this literature review and its main contribution."
            sections[step] = generate_section(llm, step, instr, cfg.topic, papers, per_section)
        else:
            theme_obj = next((t for t in themes if t.get("name") == step), None)
            sections[step] = generate_theme_section(llm, theme_obj or {}, cfg.topic, papers_by_key, per_section) if theme_obj else ""

    return sections


# =========================================================================================
# Phase 13: Professional document formatting
# =========================================================================================

def _set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    section = doc.sections[0]
    section.page_height = Cm(29.7)   # A4
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def add_hyperlink(paragraph, url: str, text: str, font_name: str = "Times New Roman",
                   font_size: int = 12, color: str = "1155CC", underline: bool = True):
    """python-docx has no native hyperlink support - this builds the required raw
    <w:hyperlink> XML by hand so citations are real, clickable links in Word (not just
    styled text), pointing straight at the paper's arXiv page."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rpr.append(rfonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size * 2))
    rpr.append(sz)

    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rpr.append(c)

    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rpr.append(u)

    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_plain_run(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def _new_body_paragraph(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    return p


CITATION_TAG_RE = re.compile(r"\[(P\d+)\]")


def _add_body_paragraph_with_citations(doc: Document, text: str, key_to_paper: Dict[str, PaperRecord],
                                        style: str, ieee_order: List[str]) -> None:
    """Every [P3]-style citation tag becomes a real, clickable Word hyperlink pointing directly
    at that paper's arXiv page, formatted for the chosen citation style (e.g. "[3]" for IEEE,
    "(Smith, 2023)" for APA/MLA/Chicago)."""
    for para_text in text.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = _new_body_paragraph(doc)
        pos = 0
        for match in CITATION_TAG_RE.finditer(para_text):
            if match.start() > pos:
                _add_plain_run(p, para_text[pos:match.start()])
            key = match.group(1)
            paper = key_to_paper.get(key)
            if paper:
                if style == "IEEE":
                    if key not in ieee_order:
                        ieee_order.append(key)
                    label = f"[{ieee_order.index(key) + 1}]"
                else:
                    label = f"({paper.author_display()}, {paper.year})"
                url = f"https://arxiv.org/abs/{paper.arxiv_id}"
                add_hyperlink(p, url, label)
            pos = match.end()
        if pos < len(para_text):
            _add_plain_run(p, para_text[pos:])


def _add_page_numbers(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)


def _add_toc_field(doc: Document) -> None:
    """Inserts a real Word TOC field. It auto-populates when the user opens the file in
    Microsoft Word and presses F9 / 'Update Field' (Word builds TOCs client-side; python-docx
    cannot pre-render page numbers without Word itself)."""
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar"); fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t"); placeholder.text = "Right-click and choose 'Update Field' to build the Table of Contents."
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_separate)
    run._r.append(placeholder); run._r.append(fld_end)
    doc.add_page_break()


def _add_reference_entry(doc: Document, p: PaperRecord, style: str,
                          ieee_number: Optional[int] = None) -> None:
    """Adds one reference-list paragraph with the trailing URL as a real, clickable
    hyperlink (a plain-text URL typed programmatically into a Word run does NOT
    auto-linkify the way it does when a human types it live in Word)."""
    full_text = format_reference_entry(p, style, ieee_number=ieee_number)
    url = f"https://arxiv.org/abs/{p.arxiv_id}"
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    if url in full_text:
        before = full_text.split(url)[0]
        _add_plain_run(para, before)
        add_hyperlink(para, url, url)
    else:
        _add_plain_run(para, full_text)


def build_review_docx(cfg: ReviewConfig, sections: Dict[str, str], comparison_df: pd.DataFrame,
                       papers: List[PaperRecord], gaps: List[Dict[str, Any]],
                       ieee_order: List[str]) -> str:
    doc = Document()
    _set_base_style(doc)
    _add_page_numbers(doc)

    title = doc.add_heading(f"A Literature Review on {cfg.topic}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Generated {datetime.now().strftime('%B %d, %Y')} | Citation style: {cfg.citation_style}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    _add_toc_field(doc)

    key_to_paper = {p.citation_key: p for p in papers}

    for heading, body in sections.items():
        if not body:
            continue
        doc.add_heading(heading, level=1)
        _add_body_paragraph_with_citations(doc, body, key_to_paper, cfg.citation_style, ieee_order)

    doc.add_heading("Comparative Summary of Reviewed Literature", level=1)
    if not comparison_df.empty:
        cols = [c for c in comparison_df.columns if c != "Key"]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Light Grid Accent 1"
        for i, c in enumerate(cols):
            table.rows[0].cells[i].text = c
        for _, r in comparison_df.iterrows():
            cells = table.add_row().cells
            for i, c in enumerate(cols):
                cells[i].text = str(r[c])[:300]

    doc.add_heading("References", level=1)
    if cfg.citation_style == "IEEE":
        ordered = [key_to_paper[k] for k in ieee_order if k in key_to_paper]
        remaining = [p for p in papers if p.citation_key not in ieee_order]
        for p in remaining:
            ieee_order.append(p.citation_key)
        ordered = ordered + remaining
        for i, p in enumerate(ordered, start=1):
            _add_reference_entry(doc, p, "IEEE", ieee_number=i)
    else:
        ordered = sorted(papers, key=lambda p: p.author_display())
        for p in ordered:
            _add_reference_entry(doc, p, cfg.citation_style)

    path = "/tmp/Literature_Review.docx"
    doc.save(path)
    return path


# =========================================================================================
# Phase 14: Export options (BibTeX / CSV)
# =========================================================================================

def build_bibtex(papers: List[PaperRecord]) -> str:
    entries = []
    for p in papers:
        first_author_last = p.authors[0].split()[-1] if p.authors else "unknown"
        cite_id = f"{first_author_last}{p.year}{p.arxiv_id[-4:]}".replace(".", "")
        authors_bib = " and ".join(p.authors)
        entries.append(
            f"@article{{{cite_id},\n"
            f"  title={{{p.title}}},\n"
            f"  author={{{authors_bib}}},\n"
            f"  journal={{arXiv preprint arXiv:{p.arxiv_id}}},\n"
            f"  year={{{p.year}}},\n"
            f"  url={{https://arxiv.org/abs/{p.arxiv_id}}}\n"
            f"}}"
        )
    return "\n\n".join(entries)


def build_papers_csv(papers: List[PaperRecord]) -> pd.DataFrame:
    rows = []
    for p in papers:
        rows.append({
            "Citation Key": p.citation_key,
            "Title": p.title,
            "Authors": "; ".join(p.authors),
            "Year": p.year,
            "arXiv ID": p.arxiv_id,
            "PDF URL": p.pdf_url,
            "Relevance Score": p.relevance_score,
            "Themes": "; ".join(p.themes),
            "PDF Extracted OK": p.pdf_ok,
        })
    return pd.DataFrame(rows)


# =========================================================================================
# Full pipeline orchestrator
# =========================================================================================

def run_full_pipeline(api_key: str, model_name: str, cfg: ReviewConfig, progress=None):
    """Runs every phase of the agent end-to-end and returns all artifacts needed by the UI."""
    llm = LLMClient(api_key=(api_key or "").strip(), model=model_name or MODEL_NAME_DEFAULT)

    papers, tried_queries = run_search_and_select(llm, cfg, progress=progress)
    if not papers:
        raise RuntimeError("No papers were found for this topic. Try a broader or differently worded topic.")

    download_and_extract(papers, progress=progress)
    summarize_all_papers(llm, cfg, papers, progress=progress)

    if progress:
        progress(0.68, desc="Discovering research themes...")
    themes = discover_themes(llm, cfg.topic, papers)

    if progress:
        progress(0.70, desc="Detecting research gaps...")
    gaps = detect_gaps(llm, cfg.topic, papers, themes)

    if progress:
        progress(0.72, desc="Detecting contradictory findings...")
    contradictions = detect_contradictions(llm, papers)

    sections = build_full_review_text(llm, cfg, papers, themes, gaps, contradictions, progress=progress)

    if progress:
        progress(0.97, desc="Formatting document...")
    comparison_df = build_comparison_table(papers)
    ieee_order: List[str] = []
    docx_path = build_review_docx(cfg, sections, comparison_df, papers, gaps, ieee_order)
    bibtex_str = build_bibtex(papers)
    papers_df = build_papers_csv(papers)

    if progress:
        progress(1.0, desc="Done.")

    return {
        "papers": papers,
        "themes": themes,
        "gaps": gaps,
        "contradictions": contradictions,
        "sections": sections,
        "comparison_df": comparison_df,
        "papers_df": papers_df,
        "docx_path": docx_path,
        "bibtex_str": bibtex_str,
        "tried_queries": tried_queries,
    }


# =========================================================================================
# STREAMLIT UI
# =========================================================================================

st.set_page_config(
    page_title="AI Literature Review Agent",
    page_icon="📚",
    layout="wide",
)

CUSTOM_CSS = """
<style>
.app-header {
    background: linear-gradient(135deg, #161f38 0%, #1e2a4d 60%, #241a4d 100%);
    border: 1px solid #2a3a5c;
    border-radius: 14px;
    padding: 28px 32px;
    margin-bottom: 18px;
}
.app-header h1 {
    font-size: 1.9rem; font-weight: 700; letter-spacing: -0.02em; margin-bottom: 4px;
    background: linear-gradient(90deg,#a5b4fc,#e9d5ff);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
}
.app-header p {color: #9fb3d9; font-size: 0.95rem; margin: 0;}
.badge-row {display:flex; gap:8px; margin-top:14px; flex-wrap:wrap;}
.badge {
    font-size: 0.72rem; font-weight: 600; letter-spacing:.03em; text-transform:uppercase;
    color:#c7d2fe; background:#1e2a4d; border:1px solid #33427a;
    padding:4px 10px; border-radius:999px;
}
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

st.markdown(
    """
    <div class="app-header">
      <h1>AI-Powered Autonomous Literature Review Agent</h1>
      <p>Plans a search strategy, autonomously searches arXiv, evaluates coverage and
      replans if needed, clusters themes, detects gaps and contradictions, and writes a
      fully cited, professionally formatted literature review.</p>
      <div class="badge-row">
        <span class="badge">Autonomous Search</span>
        <span class="badge">Coverage-Based Replanning</span>
        <span class="badge">Theme Clustering</span>
        <span class="badge">Gap Detection</span>
        <span class="badge">Auto-Cited Word Export</span>
      </div>
    </div>
    """,
    unsafe_allow_html=True,
)

# --- Sidebar: API settings ---
with st.sidebar:
    st.header("🔑 API Settings")
    api_key = st.text_input(
        "Groq API Key", type="password", placeholder="gsk_...",
        help="Never stored - used only for this session. Get one at https://console.groq.com/keys",
    )
    model_name = st.text_input("Model", value=MODEL_NAME_DEFAULT)

# --- Main inputs ---
col1, col2 = st.columns([2, 1])
with col1:
    topic = st.text_input(
        "Research Topic / Question",
        placeholder="e.g. Deep Learning for Electricity Demand Forecasting",
    )
with col2:
    num_papers = st.slider("Number of Papers to Analyze", 5, 25, 10, step=1)

col3, col4, col5, col6 = st.columns(4)
with col3:
    depth = st.selectbox("Review Depth (search rounds & replanning)", ["Quick", "Standard", "Comprehensive"], index=1)
with col4:
    length = st.selectbox("Desired Review Length", ["3-5 pages", "5-10 pages", "10-20 pages", "Custom"], index=1)
with col5:
    custom_words = st.number_input("Custom word count (used only if 'Custom')", value=3000, step=250)
with col6:
    citation_style = st.selectbox("Citation Style", ["IEEE", "APA", "MLA", "Chicago"], index=0)

generate_clicked = st.button("🚀 Generate Literature Review", type="primary")

if generate_clicked:
    if not topic or not topic.strip():
        st.warning("⚠️ Please enter a research topic.")
    elif not api_key or not api_key.strip():
        st.warning("⚠️ Please enter your Groq API key in the sidebar.")
    else:
        cfg = ReviewConfig(
            topic=topic.strip(),
            num_papers=int(num_papers),
            depth=depth,
            length=length,
            custom_words=int(custom_words) if custom_words else 3000,
            citation_style=citation_style,
        )

        progress_bar = st.progress(0)
        status_text = st.empty()

        def progress_cb(frac, desc=""):
            progress_bar.progress(min(max(int(frac * 100), 0), 100))
            if desc:
                status_text.text(desc)

        try:
            with st.spinner("Running the agent - this can take a few minutes..."):
                result = run_full_pipeline(api_key, model_name, cfg, progress=progress_cb)
            st.session_state["result"] = result
            st.session_state["result_cfg"] = cfg
            status_text.empty()
            progress_bar.empty()
            st.success(f"✅ Reviewed {len(result['papers'])} papers across {len(result['themes'])} themes.")
        except Exception as e:
            status_text.empty()
            progress_bar.empty()
            msg = str(e)
            if "RateLimitError" in msg or "rate_limit" in msg.lower() or "429" in msg:
                st.error(
                    "❌ Groq rate limit reached. The agent already retries and slows itself down "
                    "automatically, but your account's per-minute quota was exceeded anyway. "
                    "Try: (1) waiting ~60 seconds and running again, (2) lowering 'Number of papers' "
                    "or using 'Quick' depth to make fewer LLM calls, or (3) checking your usage/limits "
                    "at https://console.groq.com/settings/limits."
                )
            else:
                st.error(f"❌ Error: {e}")
                with st.expander("Show traceback"):
                    st.code(traceback.format_exc())

# --- Results (persist across reruns via session_state) ---
if "result" in st.session_state:
    result = st.session_state["result"]
    cfg = st.session_state["result_cfg"]

    tab1, tab2, tab3, tab4 = st.tabs(
        ["📄 Review Preview", "📊 Comparative Table", "🔍 Research Gaps & Contradictions", "⬇️ Downloads"]
    )

    with tab1:
        preview_md = f"# A Literature Review on {cfg.topic}\n\n"
        for heading, body in result["sections"].items():
            if body:
                preview_md += f"## {heading}\n\n{body}\n\n"
        st.markdown(preview_md)

    with tab2:
        st.dataframe(result["comparison_df"], use_container_width=True)

    with tab3:
        gaps_text = "\n\n".join(
            f"**Gap:** {g.get('gap')}\n\n*Evidence:* {g.get('evidence_keys')}"
            for g in result["gaps"]
        ) or "No specific gaps identified."
        if result["contradictions"]:
            gaps_text += "\n\n---\n\n**Conflicting Findings:**\n\n" + "\n\n".join(
                f"- {c.get('description')} (papers: {c.get('paper_keys')}) — possible reason: {c.get('possible_reason')}"
                for c in result["contradictions"]
            )
        st.markdown(gaps_text)

    with tab4:
        st.caption(
            "Word document is fully formatted (Times New Roman, justified, 1.5 spacing, "
            "headings, comparison table, references). Open in Word and right-click the "
            "Table of Contents to 'Update Field' to populate it."
        )
        with open(result["docx_path"], "rb") as f:
            st.download_button(
                "Download Literature Review (.docx)", f,
                file_name="Literature_Review.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        st.download_button(
            "Download References (.bib)", result["bibtex_str"],
            file_name="references.bib", mime="text/plain",
        )
        csv_bytes = result["papers_df"].to_csv(index=False).encode("utf-8")
        st.download_button(
            "Download Analyzed Papers (.csv)", csv_bytes,
            file_name="analyzed_papers.csv", mime="text/csv",
        )

st.markdown("---")
st.caption(
    "This agent grounds its writing in retrieved paper summaries, but always verify "
    "citations and claims against the original papers before submission."
)
